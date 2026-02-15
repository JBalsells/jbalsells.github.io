#!/usr/bin/env python3
"""Generate jorge-balsells-cv.docx — compact 2-page CV with teal palette."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Color palette ─────────────────────────────────────────────────────────
PRIMARY       = RGBColor(0x1A, 0x4F, 0x64)
PRIMARY_MED   = RGBColor(0x2A, 0x7A, 0x96)
PRIMARY_DARK  = RGBColor(0x12, 0x3A, 0x4A)
PRIMARY_LIGHT = "E8F1F5"
TEXT          = RGBColor(0x33, 0x33, 0x33)
MUTED         = RGBColor(0x6C, 0x75, 0x7D)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

os.chdir(os.path.dirname(os.path.abspath(__file__)))
doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin = Cm(1.0)
    s.bottom_margin = Cm(0.8)
    s.left_margin = Cm(1.5)
    s.right_margin = Cm(1.5)

ns = doc.styles['Normal']
ns.font.name = 'Calibri'
ns.font.size = Pt(9)
ns.font.color.rgb = TEXT
ns.paragraph_format.space_after = Pt(0)
ns.paragraph_format.space_before = Pt(0)
ns.paragraph_format.line_spacing = 1.0


# ── Helpers ────────────────────────────────────────────────────────────────

def shade(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'))

def cell_borders(cell, left="nil", top="nil", right="nil", bottom="nil"):
    """Set individual cell borders. Use 'nil' to remove or 'single|sz|color' for a border."""
    parts = []
    for side, val in [("left",left),("top",top),("right",right),("bottom",bottom)]:
        if val == "nil":
            parts.append(f'<w:{side} w:val="nil"/>')
        else:
            parts.append(f'<w:{side} {val}/>')
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>{"".join(parts)}</w:tcBorders>'))

def cell_margins(cell, top=40, bottom=40, left=80, right=60):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'))

def cell_width(cell, width_dxa):
    cell._tc.get_or_add_tcPr().append(parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{width_dxa}" w:type="dxa"/>'))

def full_width(tbl):
    pr = tbl._tbl.tblPr if tbl._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    pr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>'))
    pr.append(parse_xml(f'<w:tblCellMar {nsdecls("w")}>'
        f'<w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/>'
        f'</w:tblCellMar>'))

def no_spacing(tbl):
    """Remove spacing between table cells."""
    pr = tbl._tbl.tblPr if tbl._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    pr.append(parse_xml(f'<w:tblCellSpacing {nsdecls("w")} w:w="0" w:type="dxa"/>'))

def run(paragraph, text, size=9, color=TEXT, bold=False):
    r = paragraph.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.name = 'Calibri'
    return r


# ── Section heading ───────────────────────────────────────────────────────

def section_bar(text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    full_width(tbl)
    cell = tbl.rows[0].cells[0]
    shade(cell, "1A4F64")
    cell_borders(cell, bottom='w:val="single" w:sz="6" w:color="2A7A96" w:space="0"')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run(p, f"  {text}", size=9.5, color=WHITE, bold=True)


# ── Two-column entry (period sidebar + content) ──────────────────────────

def entry(title, org, period, bullets=None, note=None):
    """Two-column entry: teal period sidebar | content area."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    full_width(tbl)
    no_spacing(tbl)

    # Left column: period
    lc = tbl.rows[0].cells[0]
    cell_width(lc, 1300)  # ~2.3cm
    shade(lc, "1A4F64")
    cell_borders(lc, right='w:val="single" w:sz="6" w:color="2A7A96" w:space="0"')
    cell_margins(lc, top=40, bottom=40, left=50, right=50)
    pl = lc.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.paragraph_format.space_before = Pt(2)
    pl.paragraph_format.space_after = Pt(2)
    # Split period on " – " only for two-line display
    parts = [p.strip() for p in period.split("–")] if "–" in period else [period]
    for i, line in enumerate(parts):
        if i > 0:
            pl.add_run("\n").font.size = Pt(2)
        run(pl, line.strip(), size=7, color=WHITE, bold=True)

    # Right column: content
    rc = tbl.rows[0].cells[1]
    shade(rc, PRIMARY_LIGHT)
    cell_borders(rc, left='w:val="single" w:sz="6" w:color="2A7A96" w:space="0"')
    cell_margins(rc, top=35, bottom=35, left=90, right=50)

    # Title + org
    pt = rc.paragraphs[0]
    pt.paragraph_format.space_after = Pt(1)
    run(pt, title, size=9, color=PRIMARY_DARK, bold=True)
    if org:
        run(pt, f"  —  {org}", size=8, color=MUTED)

    # Note (GPA etc)
    if note:
        pn = rc.add_paragraph()
        pn.paragraph_format.space_after = Pt(0)
        run(pn, note, size=7.5, color=PRIMARY_MED, bold=True)

    # Bullets
    if bullets:
        for b in bullets:
            pb = rc.add_paragraph()
            pb.paragraph_format.space_after = Pt(0)
            pb.paragraph_format.left_indent = Cm(0.25)
            run(pb, "▸ ", size=6.5, color=PRIMARY_MED)
            run(pb, b, size=8, color=TEXT)

    # Tiny spacer
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(1)
    sp.paragraph_format.space_after = Pt(0)


# ── Compact single-line entry ─────────────────────────────────────────────

def compact_entry(title, org, period, desc=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.2)
    run(p, "▸ ", size=7, color=PRIMARY_MED)
    run(p, title, size=8.5, color=PRIMARY_DARK, bold=True)
    if org:
        run(p, f"  —  {org}", size=8, color=MUTED)
    run(p, f"  [{period}]", size=7.5, color=PRIMARY)
    if desc:
        dp = doc.add_paragraph()
        dp.paragraph_format.space_before = Pt(0)
        dp.paragraph_format.space_after = Pt(1)
        dp.paragraph_format.left_indent = Cm(0.5)
        run(dp, desc, size=7.5, color=MUTED)


# ═══════════════════════════════════════════════════════════════════════════
#  HEADER
# ═══════════════════════════════════════════════════════════════════════════

htbl = doc.add_table(rows=1, cols=1)
htbl.alignment = WD_TABLE_ALIGNMENT.CENTER
full_width(htbl)
hc = htbl.rows[0].cells[0]
shade(hc, "123A4A")
cell_borders(hc)

p = hc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(1)
run(p, "JORGE A. BALSELLS ORELLANA", size=18, color=WHITE, bold=True)

p2 = hc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(1)
run(p2, "Electronic Engineer  |  Master's in Mathematical Engineering  |  Master's in AI",
    size=9.5, color=RGBColor(0xCC,0xDD,0xE6))

p3 = hc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(8)
run(p3, "10+ years firmware & software  ·  4+ years MLOps  ·  14 years volunteer firefighter",
    size=8, color=RGBColor(0xAA,0xBB,0xCC))

# Contact bar
ctbl = doc.add_table(rows=1, cols=5)
ctbl.alignment = WD_TABLE_ALIGNMENT.CENTER
full_width(ctbl)
contacts = [
    "jorgebalsells@gmail.com", "+(502) 3016-9594", "Guatemala City (GMT-6)",
    "linkedin.com/in/jbalsells", "github.com/JBalsellsO",
]
for i, txt in enumerate(contacts):
    c = ctbl.rows[0].cells[i]
    shade(c, "1A4F64")
    cell_borders(c)
    pp = c.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(4)
    pp.paragraph_format.space_after = Pt(4)
    run(pp, txt, size=7.5, color=WHITE)

doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ═══════════════════════════════════════════════════════════════════════════
#  PROFILE
# ═══════════════════════════════════════════════════════════════════════════

section_bar("PROFILE")
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(3)
run(p,
    "Electronic Engineer with Master's degrees in Mathematical Engineering and Artificial Intelligence. "
    "10+ years bridging hardware and software — from embedded systems and signal processing to scalable "
    "ML infrastructure in production. Currently a professor at ECFM-USAC (Simulation & Instrumentation labs) "
    "and development engineer for the LAGO project, while serving as MLOps Engineer at Vana.gt for the "
    "Data Science Risk team. Volunteer firefighter with SPRAT Level 1 rope access certification.",
    size=8.5, color=TEXT)

# Key highlights
hl = doc.add_table(rows=1, cols=4)
hl.alignment = WD_TABLE_ALIGNMENT.CENTER
full_width(hl)
for i, (n, l) in enumerate([("10+","YRS EXPERIENCE"),("3","MASTER'S DEGREES"),("9+","CERTIFICATES"),("14","YRS FIREFIGHTER")]):
    c = hl.rows[0].cells[i]
    shade(c, PRIMARY_LIGHT)
    cell_borders(c)
    pp = c.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pp.paragraph_format.space_before = Pt(4)
    pp.paragraph_format.space_after = Pt(4)
    run(pp, n, size=12, color=PRIMARY, bold=True)
    run(pp, f"  {l}", size=6, color=MUTED)

doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ═══════════════════════════════════════════════════════════════════════════
#  SKILLS
# ═══════════════════════════════════════════════════════════════════════════

section_bar("PROFESSIONAL SKILLS")

skills_data = [
    ("Software", ["Full Stack","APIs","Git","CI/CD","Databases","Algorithms","MLOps","Bioinformatics"]),
    ("Cloud & Infra", ["AWS","Docker","Linux","DevOps","DataOps","IaC"]),
    ("Hardware", ["Digital Electronics","Firmware","3D Prototyping","Robotics","Instrumentation"]),
    ("Academic", ["Simulation","Electronics","Programming","Research Mentoring"]),
    ("Languages", ["Spanish (Native)","English (Intermediate)"]),
]
for cat, tags in skills_data:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run(p, f"{cat.upper()}: ", size=7.5, color=PRIMARY_DARK, bold=True)
    for j, t in enumerate(tags):
        r = run(p, f" {t} ", size=8, color=PRIMARY_DARK)
        r._r.get_or_add_rPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{PRIMARY_LIGHT}" w:val="clear"/>'))
        if j < len(tags)-1:
            run(p, " ", size=7)

doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ═══════════════════════════════════════════════════════════════════════════
#  WORK EXPERIENCE
# ═══════════════════════════════════════════════════════════════════════════

section_bar("WORK EXPERIENCE")

entry("Profesor", "ECFM, Universidad de San Carlos de Guatemala", "Ene 2024 – Actualidad", [
    "Ingeniero de desarrollo en proyecto LAGO (Latin American Giant Observatory).",
    "Catedrático de Laboratorio de Simulación y Laboratorio de Instrumentación.",
])

entry("Machine Learning Engineer LII", "Vana", "2022 – Actualidad", [
    "MLOps para el equipo Data Science RISK — entrenamiento, despliegue y monitoreo de modelos a escala.",
])

entry("Profesor", "Escuela de Estudios de Postgrado, Fac. Ingeniería, USAC", "2022", [
    "Catedrático de \"Software y bases de datos biomédicas\". Evaluaciones docentes: 93/100 y 97/100.",
])

entry("Coordinador — Profesional de Laboratorio", "Fab-LAB, Dir. Gral. de Investigación, USAC", "Oct 2019 – Ene 2024", [
    "Planificación y asesoría de proyectos. Desarrollo de prototipos de innovación. Operación de CNC.",
])

entry("Full Stack Developer", "O+M Plus", "2016 – 2020", [
    "Desarrollo web y algoritmos para coberturas Wireless (Fresnel, LTE, IPRAN, GPON).",
    "APIs de Google Maps para factibilidades, reportería, depuración de BD y backups SFTP.",
])

entry("Auxiliar de Cátedra II", "Lab. Electrónica, Ingeniería, USAC", "2015 – 2017", [
    "Laboratorios de Robótica y Comunicaciones 2.",
])

entry("Desarrollador de Ingenierías", "PROSIMA", "2013 – 2014", [
    "Ingenierías de automatización industrial para TIGSA y CEMPRO.",
])

# ═══════════════════════════════════════════════════════════════════════════
#  EDUCATION
# ═══════════════════════════════════════════════════════════════════════════

section_bar("EDUCATION")

entry("Máster en Inteligencia Artificial", "CEUPE", "2023 – 2024", [
    "Deep learning, NLP y aplicaciones empresariales de IA.",
])

entry("Máster en Ingeniería Matemática y Computación", "UNIR, España", "2020 – 2021",
    ["Geometría Diferencial, Análisis Numérico, Procesamiento de Señales, Sistemas Dinámicos."],
    note="GPA: 8.68/10")

entry("Esp. en Big Data e Inteligencia Artificial", "UNIR, España", "2021",
    note="GPA: 9.5/10")

entry("Esp. en Bioinformática y Biocomputación Molecular", "USAC, Guatemala", "2019",
    note="GPA: 9.4/10")

entry("Ingeniería Electrónica", "Universidad de San Carlos de Guatemala", "2008 – 2017",
    ["Colegiado Activo No. 17029. Miembro IEEE 2010–2024."],
    note="GPA: 6.9/10")

# ═══════════════════════════════════════════════════════════════════════════
#  CERTIFICATES & COURSES
# ═══════════════════════════════════════════════════════════════════════════

section_bar("CERTIFICATES & COURSES")

compact_entry('Course: "Machine Learning in Production"', "Coursera", "2026")
compact_entry('Workshops: Scientific Instrumentation (ICTP Guatemala + ICTP-IAEA Trieste)', None, "2025")
compact_entry('Specialization: "Machine Learning" (3 courses)', "Coursera", "2025")
compact_entry('Courses: AWS Cloud Essentials · DevOps, DataOps, MLOps · Web Apps & CLI for Data Eng.', "Coursera", "2025")
compact_entry('Specialization: "Modern Robotics" (5 courses)', "Coursera", "2021–2023")
compact_entry('Specialization: "Data Structures and Algorithms" (6 courses)', "Coursera", "2018–2019")

doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ═══════════════════════════════════════════════════════════════════════════
#  PROJECTS & ACHIEVEMENTS
# ═══════════════════════════════════════════════════════════════════════════

section_bar("PROJECTS & ACHIEVEMENTS")

compact_entry("Robot multifuncional orientado al área agrícola", "Proyecto DIGI, USAC", "2021–2022")
compact_entry("Dispositivo electromecánico para Telescopio — seguimiento astronómico", "TFM, UNIR", "2020–2021")
compact_entry("Publicación: Dispositivo de asistencia ventilatoria de lazo cerrado", "Rev. CyTS Vol.7 No.3, DIGI-USAC", "2020")
compact_entry("ICLR Workshop: Classification of Plasmodium Vivax in Blood Smears", "Addis Ababa, Ethiopia", "2020")
compact_entry("Detección de fases de Plasmodium Vivax con CNNs", "TFE, USAC", "2019")
compact_entry("Mercury Robotics Challenge — 1er lugar Judge's Choice (2016 & 2017)", "CUN, Colombia", "2016–2017")
compact_entry("NASA SpaceApps Challenge — Top 5 mundial \"Best Mission Concept\"", "Guatemala", "2014")

doc.add_paragraph().paragraph_format.space_after = Pt(1)

# ═══════════════════════════════════════════════════════════════════════════
#  VOLUNTEERING & EXTRACURRICULAR (expanded)
# ═══════════════════════════════════════════════════════════════════════════

section_bar("VOLUNTEERING & EXTRACURRICULAR")

entry("Galonista II — Bombero Voluntario",
      "Benemérito Cuerpo de Bomberos Voluntarios de Guatemala",
      "2011 – Actualidad", [
    "Décima compañía, área metropolitana. 14 años de servicio activo.",
    "Firefighter III (Reds Team, FAG). Fire Instructor NFPA 1403 (Texas A&M).",
    "Control de incendios GLP (Zeta Gas). CBSCI (USAID/CVB). PRIMAP (USAID/CVB). TUM básico (CVB).",
    "Incendios forestales TTBCIF (CONRED). Estancias técnicas en Baltar y Macedo de Cavaleiros, Portugal.",
])

entry("Técnico en accesos con cuerdas — Nivel 1",
      "Society of Professional Rope Access Technicians (SPRAT)",
      "2025", [
    "Acceso con cuerdas para trabajo seguro en alturas, normativas de seguridad, técnicas de rescate y equipos especializados.",
])

entry("Técnico en rescate agreste — Nivel 2",
      "ASONBOMD — Santa Isabel",
      "2023 – 2024", [
    "Rescate agreste mediante progresión con cuerdas en entornos naturales. Evaluación de riesgos y equipos técnicos.",
])

entry("Escalada de árboles — Tree Climbing",
      "Tree Climbing — Purulhá B.V.",
      "2023", [
    "Aliado Oficial. Técnicas avanzadas de progresión con cuerdas en árboles.",
])

# ═══════════════════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════════════════

doc.save("jorge-balsells-cv.docx")
print("OK jorge-balsells-cv.docx")
